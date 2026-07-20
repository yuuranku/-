from copy import deepcopy
from datetime import datetime
from pathlib import Path
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


PICS = Path.home() / "Pictures"
DOCX = next(p for p in PICS.glob("*.docx") if p.name.startswith("1952_HZ-6"))
BACKUP = DOCX.with_name(DOCX.stem + "_backup_" + datetime.now().strftime("%Y%m%d-%H%M%S") + DOCX.suffix)

EAST_FONT = "SimSun"
LATIN_FONT = "Times New Roman"
INK = RGBColor(28, 28, 28)


def set_run_font(run, size=10.5, bold=None, color=INK):
    run.font.name = LATIN_FONT
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), EAST_FONT)
    r_fonts.set(qn("w:cs"), LATIN_FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_paragraph_text(paragraph, text, size=10.5, bold=False):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15


def insert_after(paragraph, text=None, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph._parent.add_paragraph()
    p._p = new_p
    p._element = new_p
    if style is not None:
        p.style = style
    if text is not None:
        set_paragraph_text(p, text)
    return p


def insert_heading_after(paragraph, text, style):
    p = insert_after(paragraph, style=style)
    set_paragraph_text(p, text, size=14, bold=True)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    return p


def replace_in_paragraphs(doc, mapping):
    for paragraph in doc.paragraphs:
        text = paragraph.text
        new_text = text
        for old, new in mapping.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            size = paragraph.runs[0].font.size.pt if paragraph.runs and paragraph.runs[0].font.size else 10.5
            bold = any(run.bold for run in paragraph.runs)
            set_paragraph_text(paragraph, new_text, size=size, bold=bold)


def replace_in_tables(doc, mapping):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    new_text = text
                    for old, new in mapping.items():
                        new_text = new_text.replace(old, new)
                    if new_text != text:
                        size = paragraph.runs[0].font.size.pt if paragraph.runs and paragraph.runs[0].font.size else 9.5
                        bold = any(run.bold for run in paragraph.runs)
                        set_paragraph_text(paragraph, new_text, size=size, bold=bold)


def clear_numbering(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def add_numbered_after(paragraph, text, style=None):
    p = insert_after(paragraph, style=style)
    set_paragraph_text(p, text)
    p.style = "List Number" if "List Number" in [s.name for s in paragraph.part.document.styles] else p.style
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-12)
    return p


def main():
    shutil.copy2(DOCX, BACKUP)
    doc = Document(DOCX)

    replacements = {
        "HZ-6样本线现场任务报告": "HZ-6样本现场任务报告",
        "HZ-6 样本线任务": "HZ-6 样本任务",
        "样本线现场任务": "样本现场任务",
        "低林样本线任务": "低林采样任务",
        "样本线北侧": "样本点北侧",
        "树线方向": "林缘方向",
        "转向树线": "转向林缘",
        "样本线标灯": "采样标灯",
        "离开样本线": "离开采样区域",
        "偏离样本线": "偏离采样区域",
        "路线东侧移动源": "行进区域东侧疑似移动目标",
        "队伍沿采样线前进": "队伍进入采样区域",
        "低机位逃离路线记录": "低机位逃离途中记录",
        "不无法确定是否为已知物种": "无法确定是否为已知物种",
        "哈弗大学特调研究员": "哈佛大学特调研究员",
        "美国海军借调枪": "美国海军借调；携带 M1 卡宾枪",
    }
    replace_in_paragraphs(doc, replacements)
    replace_in_tables(doc, replacements)

    paragraph_updates = {
        "任务区域位于白渊针叶层，地平线站东南测网 8.7 英里桩位，路线由 HZ-6A 车辆回转点延伸至 HZ-6C 低林样本点。HZ-6A 可由 M29C Weasel 履带车抵达；HZ-6B 至 HZ-6C 之间需步行。":
            "任务区域位于白渊针叶层，距地平线站东南测网 8.7 英里。HZ-6A 为车辆回转点，HZ-6C 为低林样本点；两点之间只能步行进入。",
        "任务按常规生态采样派出，同时补做测绘维护。派出单上列有三项：拍摄黑针木根板和地表霜晶；回收 HZ-6C 空气滤芯；更新返程标记；":
            "任务按常规生态采样派出，同时补做测绘维护。派出单简列三项：拍摄黑针木根板与地表霜晶；回收 HZ-6C 空气滤芯；更新返程标记。",
        "标准的BAS常规任务出勤小队，两名护卫前后保护研究人员摄影通讯员。":
            "按 BAS 常规出勤配置：两名护卫分列前后，保护研究人员、摄影员和无线电员。",
        "复核阶段将接触对象赋予临时编号 HZ6-CO-01。该编号只用于关联底片、足迹、伤情和现场痕迹，不无法确定是否为已知物种。":
            "复核阶段将接触对象赋予临时编号 HZ6-CO-01。该编号只用于关联底片、足迹、伤情和现场痕迹，无法确定是否为已知物种。",
        "目标贴近地面移动，长背，高肩，前肢在闪光里显得过长。它穿过黑针木根板之间的空隙时没有明显停顿。底片只能确认它曾以模糊剪影或局部湿亮轮廓进入画面，无法给出具体判断。":
            "目标贴近地面移动，背部较长，肩部位置偏高。闪光下可见前肢比例异常，但底片没有给出完整轮廓。它经过根板空隙时没有明显停顿，像熟悉那片低林。"
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in paragraph_updates:
            set_paragraph_text(paragraph, paragraph_updates[text])

    # Insert survivor inquiry, psychological assessment, and follow-up handling
    anchor = next(p for p in doc.paragraphs if p.text.strip() == "九、人员与物资处置")
    heading_style = anchor.style
    body_style = doc.paragraphs[3].style

    inserted = []
    current = anchor._element.getprevious()
    # Use the paragraph before the old section as insertion point.
    prev_para = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("目标贴近地面移动"):
            prev_para = p
            break
    if prev_para is None:
        prev_para = doc.paragraphs[doc.paragraphs.index(anchor) - 1]

    p = insert_heading_after(prev_para, "九、幸存者补充询问摘录", heading_style)
    inserted.append(p)
    additions = [
        "以下摘录来自地平线站医务室稳定观察后的补记，时间为 1952 年 11 月 19 日 0615 至 0648。马洛里已复温，左臂夹板固定。记录员注明：受询人能确认姓名、地点和相机编号，但一旦谈及队伍散开、底片冲洗或其他成员位置，回答明显变慢，手部震颤加重。",
        "问：你现在能说明小队是怎样散开的吗？",
        "答：不能。我一说到那一段，手就开始抖。不是想不起来，是那几分钟不在我这里。前面是拉斯克说回去，后面就是水。",
        "问：你是否看见拉斯克中尉倒下？",
        "答：我听见他叫基恩，或者叫我。我不能说我看见他倒下。不要让我替底片说话，我没有完整看见。",
        "问：为什么还在按快门？",
        "答：因为相机还在我手里。那时候它比我可靠。只要我还在过片，就像还有一件任务没有丢。停下来以后，我就只能听见他们在后面。",
        "问：你是否愿意查看冲洗后的影像，帮助回忆？",
        "答：不要。我醒着的时候已经会看见闪光。别把底片放到我眼前。如果它拍到了他们，我受不了；如果它什么都没拍到，我也受不了。",
        "问：你能描述接触对象的形态吗？",
        "答：每次要我说形状，我先想起克莱恩的样品箱碰到地上的声音。不是形状，是声音。然后就没有了。",
        "问：你认为自己为什么活下来？",
        "答：不知道。别写成我跑得快。我不记得自己怎么离开那里。你们找到我的时候相机还在，这件事比我活下来更让我难受。",
    ]
    for text in additions:
        p = insert_after(p, style=body_style)
        set_paragraph_text(p, text)
        if text.startswith("问：") or text.startswith("答："):
            p.paragraph_format.left_indent = Pt(18)
        inserted.append(p)

    p = insert_heading_after(p, "十、心理评估与后续处理", heading_style)
    inserted.append(p)
    assessment = [
        "医务官记录：受询人意识清楚，无持续谵妄，能辨认地平线站、Nikon S 相机和救援人员。主要异常为惊跳反应、回避性沉默、片段性记忆缺口，以及明显的幸存者负罪。提及底片、冲洗、队友姓名时，呼吸变浅，右手反复做过片动作；提及“是否看见他人死亡”时，回答中断两次。",
        "评估意见：马洛里的陈述可用于确认个人感受、相机使用状态和少量时间顺序，不宜作为完整现场复原依据。继续强行追问接触对象外形，可能扩大记忆混乱。限制级影像不得向其展示，也不得以影像内容诱导其补全记忆。",
        "后续处理：马洛里留站观察四十八小时，夜间由医务室记录睡眠和惊醒情况。所有关于事故经过的补充询问须有医务官在场，并限于任务流程、装备状态和救援定位。1952 年 12 月 4 日调离地平线站后，心理观察记录与摄影卷宗分开封存。对外说明仍按低温暴露、路线偏离和野外通讯失败处理。"
    ]
    for text in assessment:
        p = insert_after(p, style=body_style)
        set_paragraph_text(p, text)
        inserted.append(p)

    # Renumber old trailing sections.
    for paragraph in doc.paragraphs:
        txt = paragraph.text.strip()
        if txt == "九、人员与物资处置":
            set_paragraph_text(paragraph, "十一、人员与物资处置", size=14, bold=True)
        elif txt == "十、结论与建议":
            set_paragraph_text(paragraph, "十二、结论与建议", size=14, bold=True)

    # Replace the one-line recommendation block with separate, editable items.
    rec_para = next((p for p in doc.paragraphs if p.text.strip().startswith("HZ-6B 至 HZ-6C 低林段应暂停")), None)
    if rec_para:
        parent = rec_para._p.getparent()
        insert_point = rec_para
        recs = [
            "HZ-6B 至 HZ-6C 低林区域暂停非必要任务，直至路线、通讯和照明规程复核完成。",
            "所有低林采样任务不得少于五人，且前后均需配置安全人员。",
            "1600 当地时间后，未经站长和海军联络办公室双重批准，不得越过 HZ-6A。",
            "所有摄影胶卷应在封存链条下转交 BAS 影像分析部门。幸存人员不得接触限制级影像序列。",
            "HZ6-CO-01 不登记为新物种。后续文件继续使用临时接触对象编号，避免使用口语化、传闻化或描述性物种名。",
            "对外说明维持低温暴露、路线偏离及野外通讯失败口径。内部档案保留未分类动物接触、幸存者形态记忆不闭合和急性心理压力三项备注。",
        ]
        style = rec_para.style
        rec_para.clear()
        first_run = rec_para.add_run(recs[0])
        set_run_font(first_run, size=10.5)
        rec_para.style = "List Number"
        rec_para.paragraph_format.space_after = Pt(5)
        for rec in recs[1:]:
            newp = insert_after(insert_point, style=style)
            newp.style = "List Number"
            set_paragraph_text(newp, rec)
            insert_point = newp

    doc.save(DOCX)
    print("edited", DOCX)
    print("backup", BACKUP)


if __name__ == "__main__":
    main()
