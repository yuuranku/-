from pathlib import Path
import re

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "events" / "1952_HZ-6样本线任务报告_中文译本.md"
OUT_DOCX = ROOT / "HZ-6样本线任务记录" / "1952_HZ-6样本线任务报告_中文可编辑版.docx"
BADGE = ROOT / "HZ-6样本线任务记录" / "BAS徽.png"
PHOTO_DIR = ROOT / "HZ-6样本线任务记录" / "照片"
ASSET_DIR = ROOT / "tmp" / "pdfs" / "hz6_embedded_assets"

EAST_FONT = "SimSun"
LATIN_FONT = "Times New Roman"
UI_FONT = "Microsoft YaHei"
MONO_FONT = "Courier New"

INK = RGBColor(28, 28, 28)
MUTED = RGBColor(80, 80, 80)
BORDER = "777777"
LIGHT_FILL = "F6F2E9"
HEADER_FILL = "E9E4DA"
RED = RGBColor(150, 40, 35)

PHOTO_AFTER = {
    "1110": [2],
    "1240": [1],
    "1518": [6],
    "1712": [4],
    "1819": [3],
    "1824-1828": [5],
    "1829": [7, 8],
    "1831-1836": [9, 10, 11, 13],
    "1840": [14, 17],
    "1905": [12, 15],
}


def twips(inches):
    return int(inches * 1440)


def set_run_font(run, east=EAST_FONT, latin=LATIN_FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east)
    r_fonts.set(qn("w:cs"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_font(style, east=EAST_FONT, latin=LATIN_FONT, size=10.5, color=INK, bold=False):
    font = style.font
    font.name = latin
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = color
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east)
    r_fonts.set(qn("w:cs"), latin)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size="4", value="single"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), value)
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_bottom_border(cell, color="000000", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    element = borders.find(qn("w:bottom"))
    if element is None:
        element = OxmlElement("w:bottom")
        borders.append(element)
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), size)
    element.set(qn("w:space"), "0")
    element.set(qn("w:color"), color)


def set_table_fixed_width(table, width_in, col_widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(twips(width_in)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant = tr_pr.find(qn("w:cantSplit"))
    if cant is None:
        cant = OxmlElement("w:cantSplit")
        tr_pr.append(cant)


def mark_row_as_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def clear_cell(cell):
    cell._tc.clear_content()
    return cell.add_paragraph()


def add_paragraph_text(paragraph, text, size=10.5, color=INK, east=EAST_FONT, latin=LATIN_FONT):
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if bold else part
        run = paragraph.add_run(clean)
        set_run_font(run, east=east, latin=latin, size=size, color=color, bold=bold)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    add_paragraph_text(p, text)
    return p


def add_numbered_paragraph(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    add_paragraph_text(p, text)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, east=UI_FONT, latin=LATIN_FONT, size=14, color=INK, bold=True)
    return p


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = 1
    set_table_fixed_width(table, 6.95, [6.95])
    mark_row_as_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_border(cell, color="8C8578", size="4")
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, top=110, start=160, bottom=110, end=160)
    p = clear_cell(cell)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    clean_lines = list(lines)
    while clean_lines and clean_lines[0] == "":
        clean_lines.pop(0)
    while clean_lines and clean_lines[-1] == "":
        clean_lines.pop()
    for idx, line in enumerate(clean_lines):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, east=EAST_FONT, latin=MONO_FONT, size=9.2, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_markdown_table(doc, raw_lines):
    rows = []
    for line in raw_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(set(c) <= {"-", ":", " "} for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = 1
    widths = [1.35, 1.25, 1.3, 3.05] if col_count == 4 else [6.95 / col_count] * col_count
    set_table_fixed_width(table, 6.95, widths)
    mark_row_as_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_border(cell)
            set_cell_margins(cell, top=90, start=110, bottom=90, end=110)
            if r_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
            p = clear_cell(cell)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            add_paragraph_text(p, text, size=9.6, east=EAST_FONT, latin=LATIN_FONT)
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)


def photo_caption_map():
    mapping = {}
    for path in PHOTO_DIR.glob("HZ6-R06-*.png"):
        m = re.match(r"HZ6-R06-(\d+)_([^.]*)\.png", path.name)
        if m:
            mapping[int(m.group(1))] = m.group(2)
    return mapping


CAPTIONS = photo_caption_map()


def asset_for_frame(frame_num):
    return ASSET_DIR / f"zh_photo_{frame_num:02d}.jpg"


def add_photo_cell(cell, frame_num):
    set_cell_margins(cell, top=110, start=120, bottom=90, end=120)
    set_cell_border(cell, color="8C8578", size="4")
    set_cell_shading(cell, LIGHT_FILL)
    p = clear_cell(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path = asset_for_frame(frame_num)
    caption = f"HZ-6/R06-F{frame_num:02d}  {CAPTIONS.get(frame_num, '')}"
    if img_path.exists():
        with Image.open(img_path) as im:
            width_in = 2.85
            if im.height > im.width:
                width_in = 2.15
        run = p.add_run()
        shape = run.add_picture(str(img_path), width=Inches(width_in))
        shape._inline.docPr.set("title", f"HZ-6/R06-F{frame_num:02d}")
        shape._inline.docPr.set("descr", caption)
    cp = cell.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(3)
    cp.paragraph_format.space_after = Pt(0)
    cr = cp.add_run(caption)
    set_run_font(cr, east=EAST_FONT, latin=MONO_FONT, size=8.1, color=MUTED)


def add_photo_block(doc, frame_nums):
    table = doc.add_table(rows=1 + ((len(frame_nums) + 1) // 2), cols=2)
    table.alignment = 1
    set_table_fixed_width(table, 6.95, [3.475, 3.475])
    mark_row_as_header(table.rows[0])
    header = table.cell(0, 0).merge(table.cell(0, 1))
    set_cell_border(header, color="8C8578", size="4")
    set_cell_shading(header, LIGHT_FILL)
    set_cell_margins(header, top=80, start=120, bottom=80, end=120)
    p = clear_cell(header)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("影像证据并入：")
    set_run_font(r, east=EAST_FONT, latin=LATIN_FONT, size=9.2, color=MUTED)
    for row in table.rows:
        set_row_cant_split(row)
    idx = 0
    for row_idx in range(1, len(table.rows)):
        for col_idx in range(2):
            cell = table.cell(row_idx, col_idx)
            if idx < len(frame_nums):
                add_photo_cell(cell, frame_nums[idx])
            else:
                set_cell_border(cell, color="8C8578", size="4")
                set_cell_shading(cell, LIGHT_FILL)
                clear_cell(cell)
            idx += 1
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)


def add_field(paragraph, instr):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(text)
    run._r.append(fld_end)
    return run


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.28)

    styles = doc.styles
    style_font(styles["Normal"], size=10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    style_font(styles["Heading 1"], east=UI_FONT, size=14, bold=True)
    style_font(styles["List Number"], size=10.5)

    header = section.header
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)
    h_table = header.add_table(rows=1, cols=3, width=Inches(7.06))
    h_table.alignment = 1
    set_table_fixed_width(h_table, 7.06, [1.2, 4.66, 1.2])
    mark_row_as_header(h_table.rows[0])
    cells = h_table.rows[0].cells
    for cell in cells:
        set_bottom_border(cell, color="000000", size="6")
        set_cell_margins(cell, top=0, start=0, bottom=55, end=0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p0 = clear_cell(cells[0])
    if BADGE.exists():
        shape = p0.add_run().add_picture(str(BADGE), width=Inches(0.42))
        shape._inline.docPr.set("title", "BAS徽章")
        shape._inline.docPr.set("descr", "美国南极测绘局BAS徽章")
    p1 = clear_cell(cells[1])
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p1.add_run("CONFIDENTIAL  /  BAS INTERNAL")
    set_run_font(r, east=MONO_FONT, latin=MONO_FONT, size=8.8, color=INK)
    p2 = clear_cell(cells[2])
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = p2.add_run("FILE COPY")
    set_run_font(sr, east=MONO_FONT, latin=MONO_FONT, size=10, color=RED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(2)
    fr = fp.add_run("BAS-FLD-1952-HZ6-014  /  CN COPY     PAGE ")
    set_run_font(fr, east=MONO_FONT, latin=MONO_FONT, size=8.5, color=MUTED)
    page_run = add_field(fp, "PAGE")
    set_run_font(page_run, east=MONO_FONT, latin=MONO_FONT, size=8.5, color=MUTED)
    return doc


def build_docx():
    md = SOURCE_MD.read_text(encoding="utf-8")
    lines = md.splitlines()
    doc = setup_document()

    title = lines[0].lstrip("# ").strip()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(22)
    title_p.paragraph_format.space_after = Pt(18)
    tr = title_p.add_run(title)
    set_run_font(tr, east=UI_FONT, latin=LATIN_FONT, size=17.5, color=INK, bold=True)

    i = 1
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("```"):
            block = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            add_code_block(doc, block)
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip())
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        paragraph_lines = [line.strip()]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if not nxt.strip() or nxt.startswith("## ") or nxt.startswith("```") or nxt.startswith("|"):
                break
            paragraph_lines.append(nxt.strip())
            i += 1
        text = " ".join(paragraph_lines)

        number_match = re.match(r"^(\d+)\.\s+(.*)$", text)
        if number_match:
            add_numbered_paragraph(doc, number_match.group(2))
        else:
            add_body_paragraph(doc, text)

        time_match = re.match(r"^\*\*([^*]+)\*\*", text)
        if time_match:
            key = time_match.group(1)
            if key in PHOTO_AFTER:
                add_photo_block(doc, PHOTO_AFTER[key])

    doc.core_properties.title = title
    doc.core_properties.subject = "HZ-6样本线现场任务报告中文可编辑版"
    doc.core_properties.author = "BAS Horizon Station Field Office"
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build_docx())
