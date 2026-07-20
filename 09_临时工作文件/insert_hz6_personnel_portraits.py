from datetime import datetime
from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\yuuranko\Documents\白渊")
PICS = Path.home() / "Pictures"
DOCX = next(
    p for p in PICS.glob("1952_HZ-6*.docx")
    if "样本线任务报告" in p.name and "backup" not in p.name
)
PORTRAIT_DIR = ROOT / "HZ-6样本线任务记录" / "人物"
BACKUP = DOCX.with_name(DOCX.stem + "_backup_before_portraits_" + datetime.now().strftime("%Y%m%d-%H%M%S") + DOCX.suffix)

EAST_FONT = "SimSun"
LATIN_FONT = "Times New Roman"
INK = RGBColor(28, 28, 28)
BORDER = "777777"
HEADER_FILL = "E9E4DA"

PEOPLE = [
    ("霍华德·P·拉斯克中尉", "01_Howard_P_Rusk_霍华德_P_拉斯克.png", "霍华德·P·拉斯克中尉档案照"),
    ("海伦·M·克莱恩博士", "02_Helen_M_Klein_海伦_M_克莱因.png", "海伦·M·克莱恩博士档案照"),
    ("丹尼尔·基恩中士", "03_Daniel_Keene_丹尼尔_基恩.png", "丹尼尔·基恩中士档案照"),
    ("托马斯·E·马洛里", "04_Thomas_E_Mallory_托马斯_E_马洛里.png", "托马斯·E·马洛里档案照"),
    ("塞缪尔·R·万斯二级士官", "05_Samuel_R_Vance_塞缪尔_R_万斯.png", "塞缪尔·R·万斯二级士官档案照"),
]


def set_run_font(run, size=9.5, bold=False, color=INK):
    run.font.name = LATIN_FONT
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), EAST_FONT)
    r_fonts.set(qn("w:cs"), LATIN_FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def clear_cell(cell):
    cell._tc.clear_content()
    return cell.add_paragraph()


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def mark_row_as_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_table_width(table, widths):
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(total * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def main():
    shutil.copy2(DOCX, BACKUP)
    doc = Document(DOCX)
    table = doc.tables[1]

    header_text = [cell.text.strip() for cell in table.rows[0].cells]
    if "档案照" not in header_text:
        table.add_column(Inches(0.78))

    # Re-read cells after structural change.
    header = table.rows[0].cells[-1]
    p = clear_cell(header)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("档案照")
    set_run_font(r, size=9.5, bold=True)
    mark_row_as_header(table.rows[0])

    people_by_name = {name: (PORTRAIT_DIR / filename, alt) for name, filename, alt in PEOPLE}
    for row in table.rows[1:]:
        name = row.cells[0].text.strip()
        portrait_path, alt = people_by_name[name]
        cell = row.cells[-1]
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = p.add_run().add_picture(str(portrait_path), width=Inches(0.68))
        shape._inline.docPr.set("title", alt)
        shape._inline.docPr.set("descr", alt)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    widths = [1.28, 0.78, 1.05, 2.70, 0.78]
    set_table_width(table, widths)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell, top=80, start=85, bottom=80, end=85)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    if run.text:
                        set_run_font(run, size=9.2, bold=(row_idx == 0))

    doc.save(DOCX)
    print("edited", DOCX)
    print("backup", BACKUP)


if __name__ == "__main__":
    main()
