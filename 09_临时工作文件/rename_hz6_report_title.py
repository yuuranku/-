from datetime import datetime
from pathlib import Path
import shutil
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


PICS = Path.home() / "Pictures"
DOCX = next(
    p for p in PICS.glob("1952_HZ-6*.docx")
    if "\u6837\u672c\u7ebf\u4efb\u52a1\u62a5\u544a" in p.name and "backup" not in p.name
)
NEW_TITLE = "HZ-6\u91c7\u96c6\u4e0e\u56de\u6536\u4efb\u52a1\u62a5\u544a"
BACKUP = DOCX.with_name(DOCX.stem + "_backup_before_title_rename_" + datetime.now().strftime("%Y%m%d-%H%M%S") + DOCX.suffix)

EAST_FONT = "SimSun"
LATIN_FONT = "Times New Roman"
INK = RGBColor(28, 28, 28)


def set_run_font(run, east=EAST_FONT, latin=LATIN_FONT, size=None, bold=None, color=INK):
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east)
    r_fonts.set(qn("w:cs"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = color


def replace_paragraph_text(paragraph, text):
    old_runs = list(paragraph.runs)
    size = old_runs[0].font.size.pt if old_runs and old_runs[0].font.size else None
    bold = any(run.bold for run in old_runs)
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def replace_in_paragraphs(paragraphs):
    changed = 0
    for paragraph in paragraphs:
        text = paragraph.text
        if not text:
            continue
        new = text
        # Exact visible report title variants become the requested title exactly.
        for old in [
            "HZ-6\u6837\u672c\u7ebf\u73b0\u573a\u4efb\u52a1\u62a5\u544a\uff08\u4e2d\u6587\u8bd1\u672c\uff09",
            "HZ-6\u6837\u672c\u73b0\u573a\u4efb\u52a1\u62a5\u544a\uff08\u4e2d\u6587\u8bd1\u672c\uff09",
            "HZ-6\u6837\u672c\u7ebf\u73b0\u573a\u4efb\u52a1\u62a5\u544a",
            "HZ-6\u6837\u672c\u73b0\u573a\u4efb\u52a1\u62a5\u544a",
            "HZ-6\u6837\u672c\u7ebf\u4efb\u52a1\u62a5\u544a",
            "HZ-6\u6837\u672c\u4efb\u52a1\u62a5\u544a",
        ]:
            new = new.replace(old, NEW_TITLE)
        # Subject line in the first-page metadata box.
        for old in [
            "\u4e3b\u9898\uff1aHZ-6\u6837\u672c\u7ebf\u73b0\u573a\u4efb\u52a1\u53ca\u540e\u7eed\u4eba\u5458\u635f\u5931",
            "\u4e3b\u9898\uff1aHZ-6\u6837\u672c\u73b0\u573a\u4efb\u52a1\u53ca\u540e\u7eed\u4eba\u5458\u635f\u5931",
            "\u4e3b\u9898\uff1aHZ-6\u6837\u672c\u4efb\u52a1\u53ca\u540e\u7eed\u4eba\u5458\u635f\u5931",
        ]:
            new = new.replace(old, "\u4e3b\u9898\uff1a" + NEW_TITLE)
        # Last-resort cleanup if any simple residual remains in prose.
        new = new.replace("HZ-6\u6837\u672c\u7ebf", "HZ-6\u91c7\u96c6\u4e0e\u56de\u6536")
        if new != text:
            replace_paragraph_text(paragraph, new)
            changed += 1
    return changed


def walk_tables(tables):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                yield cell.paragraphs
                yield from walk_tables(cell.tables)


def all_paragraph_groups(doc):
    yield doc.paragraphs
    for section in doc.sections:
        yield section.header.paragraphs
        yield section.footer.paragraphs
        for group in walk_tables(section.header.tables):
            yield group
        for group in walk_tables(section.footer.tables):
            yield group
    for group in walk_tables(doc.tables):
        yield group


def main():
    shutil.copy2(DOCX, BACKUP)
    doc = Document(DOCX)
    changed = 0
    for paragraphs in all_paragraph_groups(doc):
        changed += replace_in_paragraphs(paragraphs)

    doc.core_properties.title = NEW_TITLE
    doc.core_properties.subject = NEW_TITLE
    doc.save(DOCX)

    with zipfile.ZipFile(DOCX) as z:
        all_text = ""
        for name in z.namelist():
            if name.startswith(("word/", "docProps/")) and name.endswith((".xml", ".rels")):
                all_text += z.read(name).decode("utf-8", errors="ignore")
        residuals = {
            "sample_line": "\u6837\u672c\u7ebf" in all_text,
            "old_title": "\u6837\u672c\u7ebf\u73b0\u573a\u4efb\u52a1\u62a5\u544a" in all_text,
            "new_title": NEW_TITLE in all_text,
        }
    print("edited", DOCX)
    print("backup", BACKUP)
    print("changed_paragraphs", changed)
    print("residuals", residuals)


if __name__ == "__main__":
    main()
