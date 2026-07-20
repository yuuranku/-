from datetime import datetime
from pathlib import Path
import shutil
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


PICS = Path.home() / "Pictures"
DOCX = next(
    p for p in PICS.glob("1952_HZ-6*.docx")
    if "backup" not in p.name
)
BACKUP = DOCX.with_name(DOCX.stem + "_backup_before_public_line_" + datetime.now().strftime("%Y%m%d-%H%M%S") + DOCX.suffix)

PUBLIC_LINE = (
    "对外仍按原口径处理：HZ-6 小队在南极科研活动中遭遇突发特大风暴，"
    "死亡及失踪人员按因公殉职处理；公开材料统一称其为极地科研任务中的牺牲者，"
    "并按规定抚恤家属。"
)

EAST_FONT = "SimSun"
LATIN_FONT = "Times New Roman"
INK = RGBColor(28, 28, 28)


def set_rfonts(run, size=None, bold=None):
    run.font.name = LATIN_FONT
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), EAST_FONT)
    r_fonts.set(qn("w:cs"), LATIN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = INK


def replace_paragraph(paragraph, text):
    old_runs = list(paragraph.runs)
    size = old_runs[0].font.size.pt if old_runs and old_runs[0].font.size else 10.5
    bold = any(run.bold for run in old_runs)
    paragraph.clear()
    run = paragraph.add_run(text)
    set_rfonts(run, size=size, bold=bold)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)


def should_replace(text):
    stripped = text.strip()
    return (
        stripped.startswith("对外仍按原口径处理")
        or stripped.startswith("对外说明")
        or "南极科研遇到突发风暴" in stripped
        or "突发性特大风暴" in stripped
        or "抚恤" in stripped
    )


def main():
    shutil.copy2(DOCX, BACKUP)
    doc = Document(DOCX)
    changed = 0

    for paragraph in doc.paragraphs:
        if should_replace(paragraph.text):
            replace_paragraph(paragraph, PUBLIC_LINE)
            changed += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if should_replace(paragraph.text):
                        replace_paragraph(paragraph, PUBLIC_LINE)
                        changed += 1

    doc.save(DOCX)

    with zipfile.ZipFile(DOCX) as z:
        text = ""
        for name in z.namelist():
            if name.startswith(("word/", "docProps/")) and name.endswith((".xml", ".rels")):
                text += z.read(name).decode("utf-8", errors="ignore")
        checks = {
            "public_line_count": text.count(PUBLIC_LINE),
            "old_short_count": text.count("南极科研遇到突发风暴"),
            "bad_yizai_count": text.count("以再南极洲"),
        }
    print("edited", DOCX)
    print("backup", BACKUP)
    print("changed", changed)
    print("checks", checks)


if __name__ == "__main__":
    main()
