from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path(r"C:\Users\yuuranko\Documents\白渊\04_任务档案\HZ-6样本线任务记录\制作参考\档案写作参考（自用）.docx")


COLORS = {
    "ink": "1E1D19",
    "muted": "5C5A52",
    "blue": "1F4D78",
    "light_blue": "E8EEF5",
    "light_gray": "F2F4F7",
    "paper": "F6F1E7",
    "fixed": "EFE7D5",
    "border": "B8B0A0",
    "red": "8A2A20",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="B8B0A0", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def set_run_font(run, name="仿宋", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph(text="", style=None, before=0, after=6, line=1.2, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        p.add_run(text)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    return p


def add_section_title(title):
    p = paragraph(style="Heading 1", before=14, after=8, line=1.15)
    run = p.add_run(title)
    set_run_font(run, "宋体", 15, COLORS["blue"], True)
    return p


def add_subtitle(title):
    p = paragraph(style="Heading 2", before=9, after=5, line=1.15)
    run = p.add_run(title)
    set_run_font(run, "宋体", 12, COLORS["blue"], True)
    return p


def add_body(text, after=6):
    p = paragraph(before=0, after=after, line=1.25)
    run = p.add_run(text)
    set_run_font(run, "仿宋", 10.5, COLORS["ink"])
    return p


def add_note(label, text, fill="F2F4F7"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    set_run_font(r, "宋体", 10.5, COLORS["blue"], True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(text)
    set_run_font(r2, "仿宋", 10, COLORS["ink"])
    paragraph("", after=2)


def add_fixed_block(title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS["fixed"])
    set_cell_border(cell, COLORS["border"])
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("固定口径｜" + title)
    set_run_font(r, "宋体", 10.5, COLORS["red"], True)
    for line in body.split("\n"):
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.line_spacing = 1.18
        r2 = p2.add_run(line)
        set_run_font(r2, "仿宋", 10, COLORS["ink"])
    paragraph("", after=2)


def add_two_col_table(headers, rows, widths=(2.2, 4.3), header_fill="E8EEF5"):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table)
    set_col_widths(table, widths)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(hdr[i], header_fill)
        set_cell_border(hdr[i])
        set_cell_margins(hdr[i])
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, "宋体", 10, COLORS["blue"], True)
    for left, right in rows:
        cells = table.add_row().cells
        for i, text in enumerate((left, right)):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cells[i])
            set_cell_margins(cells[i], top=100, bottom=100, start=140, end=140)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.18
            font = "Courier New" if any(ch.isdigit() for ch in text[:8]) or "HZ" in text or "BAS-" in text else "仿宋"
            r = p.add_run(text)
            set_run_font(r, font, 9.8, COLORS["ink"])
    paragraph("", after=3)


def add_three_col_table(headers, rows, widths=(1.65, 2.35, 2.5), header_fill="E8EEF5"):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table)
    set_col_widths(table, widths)
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        set_cell_border(cell)
        set_cell_margins(cell)
        r = cell.paragraphs[0].add_run(text)
        set_run_font(r, "宋体", 9.8, COLORS["blue"], True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_border(cells[i])
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, "仿宋", 9.5, COLORS["ink"])
    paragraph("", after=3)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.header_distance = Cm(1.2)
section.footer_distance = Cm(1.2)

styles = doc.styles
styles["Normal"].font.name = "仿宋"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
styles["Normal"].font.size = Pt(10.5)
for name in ("Heading 1", "Heading 2", "Heading 3"):
    styles[name].font.name = "宋体"
    styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hr = hp.add_run("CONFIDENTIAL / BAS INTERNAL    |    HZ-6 REFERENCE")
set_run_font(hr, "Courier New", 8.5, COLORS["muted"])

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("档案写作参考 - 自用编辑版")
set_run_font(fr, "仿宋", 8.5, COLORS["muted"])

title = paragraph(align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=4)
r = title.add_run("HZ-6 / 白渊档案写作参考")
set_run_font(r, "宋体", 22, COLORS["ink"], True)
sub = paragraph(align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=14)
r = sub.add_run("自用编辑版｜任务报告、影像说明、心理评估与封存记录")
set_run_font(r, "仿宋", 10.5, COLORS["muted"])

add_note(
    "使用说明",
    "这份文档是给你后续写白渊 / HZ-6 档案时使用的可编辑参考。固定口径、固定编号和可直接套用的段落已经单独标出；这些内容可以复制使用，除非你明确要改设定，否则建议保持原样。",
    "F2F4F7",
)

add_section_title("一、总原则")
for item in [
    ("先写事实，再写判断", "档案先记录谁在场、什么时候、在哪里、做了什么、留下了什么，最后才写哪些东西无法确认。"),
    ("异常不要说满", "可以写“未能确认”“不作为正式判断”“暂按某项记录”“无法排除”。不要直接写“怪物出现”或“它在追他们”。"),
    ("保持机构口吻", "句子要像工作人员写给上级看的，不像旁白，也不像小说独白。"),
    ("不替读者解释感受", "不要写“令人毛骨悚然”“让人不寒而栗”“这证明了……”。只写材料本身。"),
    ("新增内容要接回已有设定", "如果写新编号、新地点、新影像、新人物状态，要检查前文有没有冲突。"),
]:
    add_note(item[0], item[1], "F6F1E7")

add_section_title("二、档案语气")
add_two_col_table(
    ("可以用", "不要用"),
    [
        ("本件按……要求整理。", "恐怖的是……"),
        ("供……内部复核使用。", "可怕的一幕发生了。"),
        ("不作公开文本。", "他们终于意识到……"),
        ("暂按……记录。", "这显然说明……"),
        ("不视为正式判断。", "毫无疑问……"),
        ("无法确认是否为……", "令人震惊的是……"),
        ("未收到完整求救、伤亡报告或返程确认。", "不可名状的怪物……"),
        ("后续记录另行封存。", "真相被掩盖了。"),
        ("不得向幸存人员展示。", "它在猎杀他们。"),
        ("不宜作为完整现场复原依据。", "医生认为他见到了不可名状之物。"),
    ],
)

add_section_title("三、常用章节结构")
add_three_col_table(
    ("章节", "写什么", "注意"),
    [
        ("一、报告目的", "文件为什么存在，谁能看，不能公开到什么程度。", "保持正式，不要写宣传口吻。"),
        ("二、地点与任务概况", "地点、距离、路线点、任务目标。", "HZ-6A / HZ-6C 的位置关系不要写错。"),
        ("三、人员", "姓名、职务、状态、备注。", "不要过度写性格。"),
        ("四、主要装备与摄影器材", "装备清单、相机、胶卷、通讯设备、车辆。", "器材名称尽量固定。"),
        ("五、时间线摘要", "按时间写动作和异常。", "每条只写一件事或一组相关动作。"),
        ("六、通讯与救援", "最后通讯、求救是否完整、救援从哪里开始。", "不要新增未记录的无线电内容。"),
        ("七、影像材料", "胶卷编号、冲洗地点、帧数、归档限制。", "不要新增影像编号。"),
        ("八、接触对象", "临时编号、可见特征、无法确认的地方。", "不要命名成怪物。"),
        ("九、幸存者补充询问摘录", "可写另行封存。", "不需要写具体问答。"),
        ("十、心理评估与后续处理", "幸存者状态、记忆缺口、询问限制。", "重点写评估，不写刺激性描述。"),
        ("十一、人员与物资处置", "死亡、失踪、生还、调离、未回收物品。", "人员状态要一致。"),
        ("十二、处置意见", "路线暂停、人员配置、时间限制、影像封存、对外口径。", "固定对外口径不要随便改。"),
    ],
)

add_section_title("四、时间线写法")
add_fixed_block(
    "时间线格式",
    "0738 小队乘 M29C Weasel 从地平线站二号东门出发，沿通风车道前往 HZ-6A。\n\n0925 抵达 HZ-6A。车辆停放于标志桩外侧，基恩接入线缆中继箱。",
)
add_two_col_table(
    ("好", "不好"),
    [
        (
            "1819 基恩报告林缘方向重复啄击声停止。马洛里后述称，此时他第一次在无人命令的情况下举起相机。",
            "1819 四周突然安静下来，马洛里意识到有东西正在盯着他们。",
        )
    ],
)

add_section_title("五、异常写法")
add_three_col_table(
    ("类型", "可用写法", "不要写成"),
    [
        ("声音异常", "林缘方向存在重复啄击声。啄击声停止。无线电中出现断续语音。", "它在树林里呼唤他们。"),
        ("环境异常", "冰顶漫射光与采样标灯亮度同时下降，速度快于预期。标记点附近出现不寻常的低位拖痕。", "天气突然变得像被控制。"),
        ("影像异常", "多帧存在曝光异常。画面严重运动模糊。闪光灯并非每次触发。底片没有给出完整轮廓。", "照片揭露了怪物全貌。"),
    ],
)

add_section_title("六、接触对象写法")
add_body("接触对象不要命名成怪物。使用临时编号 HZ6-CO-01。")
add_fixed_block(
    "接触对象标准段落",
    "复核阶段将接触对象赋予临时编号 HZ6-CO-01。\n该编号只用于关联底片、足迹、伤情和现场痕迹。\n无法确定是否为已知物种。",
)
add_two_col_table(
    ("可以写", "不要写"),
    [
        ("目标贴近地面移动。", "它长着……"),
        ("背部较长。", "它故意……"),
        ("肩部位置偏高。", "它在猎杀……"),
        ("闪光下可见前肢比例异常。", "它知道他们会回来。"),
        ("底片没有给出完整轮廓。", "它在等他们。"),
    ],
)

add_section_title("七、影像材料写法")
add_fixed_block(
    "影像编号示例",
    "HZ-6/R06-F07  首次拍到未知生物残影\nHZ-6/R06-F08  更近距离接触残影\nHZ-6/R06-F09  接触后的欠曝地面废片",
)
add_fixed_block(
    "影像统计口径",
    "回收画面共三十七帧。\n六帧可供普通审查。\n十四帧列入限制归档或复核。\n十七帧为无法使用、重复、严重运动模糊或全黑画面。",
)

add_section_title("八、幸存者与心理评估写法")
add_body("幸存者不要写得像被审问。写成“补充询问”“医务官在场”“心理评估”。")
add_fixed_block(
    "心理评估标准段落",
    "原始补充询问记录另行封存，本件仅保留心理评估结论。\n\n评估意见：\n马洛里的陈述可用于确认个人感受、相机使用状态和少量时间顺序，不宜作为完整现场复原依据。\n\n继续强行追问接触对象外形，可能扩大记忆混乱。\n\n限制级影像不得向其展示，也不得以影像内容诱导其补全记忆。\n\n后续处理：\n马洛里留站观察四十八小时，夜间由医务室记录睡眠情况。",
)

add_section_title("九、对外口径写法")
add_fixed_block(
    "对外固定口径",
    "对外仍按原口径处理：HZ-6 小队在南极科研活动中遭遇突发特大风暴，死亡及失踪人员按因公殉职处理；公开材料统一称其为极地科研任务中的牺牲者，并按规定抚恤家属。",
)
add_note(
    "注意",
    "不要写“各成员殉国”，因为马洛里生还。不要把对外口径写得太像宣传稿，也不要添加英雄事迹细节，除非另写公开宣传稿。",
    "F6F1E7",
)

add_section_title("十、常用编号格式")
add_two_col_table(
    ("用途", "编号格式"),
    [
        ("任务报告", "BAS-FLD-1952-HZ6-014"),
        ("事故编号", "BAS-INC-1952-HZ6-003"),
        ("影像编号", "BAS-PHOTO-1952-HZ6-R06"),
        ("胶卷 / 帧编号", "HZ-6/R06-F07"),
        ("接触对象", "HZ6-CO-01"),
        ("医疗 / 心理记录", "BAS-MED-1952-HZ6-MAL-01"),
    ],
)
add_note(
    "编号使用限制",
    "如果任务报告里没有这个编号，正文不要随便放进去。可以在新文件里使用，但要先确定它属于新文件系统。",
    "F6F1E7",
)

add_section_title("十一、可直接套用的短段落")
fixed_paras = [
    ("报告目的", "本件按地平线站站长要求整理。内容限于 HZ-6 采集与回收任务的人员、装备、通联、影像材料和人员损失。供 BAS 内部复核、海军联络办公室查阅，以及站内野外规程修订使用。不作公开文本。"),
    ("任务区域", "任务区域位于白渊针叶层，距地平线站东南测网 8.7 英里。HZ-6A 为车辆回转点，HZ-6C 为低林样本点；两点之间只能步行进入。"),
    ("通讯", "HZ-6 低林区域无法稳定直接呼叫地平线站，任务依赖 HZ-6A 线缆中继箱。1824 至 1828 之间的断续语音是本任务最后可确认的站内接收通讯。此后未收到完整求救、伤亡报告或返程确认。"),
    ("影像限制", "马洛里未被允许查看冲洗结果。医疗理由为避免刺激和混乱记忆重建；内部备注写明：“不得允许对象根据限制级摄影序列重建记忆。”"),
    ("接触对象", "复核阶段将接触对象赋予临时编号 HZ6-CO-01。该编号只用于关联底片、足迹、伤情和现场痕迹，无法确定是否为已知物种。"),
    ("处置意见", "HZ-6B 至 HZ-6C 低林区域暂停非必要任务，直至路线、通讯和照明规程复核完成。"),
]
for title, body in fixed_paras:
    add_fixed_block(title, body)

add_section_title("十二、写完后的自查清单")
checks = [
    "有没有出现任务报告里没有的新事实？",
    "新编号是否有必要？",
    "幸存者状态是否和“马洛里生还”一致？",
    "对外口径是否仍是“南极科研活动遭遇突发特大风暴”？",
    "是否把异常说得太满？",
    "是否出现小说化、解说化、宣传化口吻？",
    "影像编号是否前后一致？",
    "HZ-6A、HZ-6B、HZ-6C 的位置关系是否写错？",
    "是否误写成“样本线”？",
    "是否把“补充询问”写成具体问答？",
]
table = doc.add_table(rows=1, cols=2)
set_table_width(table)
set_col_widths(table, (0.55, 5.95))
for i, h in enumerate((" ", "检查项")):
    cell = table.rows[0].cells[i]
    set_cell_shading(cell, COLORS["light_blue"])
    set_cell_border(cell)
    set_cell_margins(cell)
    r = cell.paragraphs[0].add_run(h)
    set_run_font(r, "宋体", 10, COLORS["blue"], True)
for idx, check in enumerate(checks, 1):
    row = table.add_row().cells
    for cell in row:
        set_cell_border(cell)
        set_cell_margins(cell)
    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = row[0].paragraphs[0].add_run(str(idx))
    set_run_font(r0, "Courier New", 9.5, COLORS["muted"])
    r1 = row[1].paragraphs[0].add_run(check)
    set_run_font(r1, "仿宋", 10, COLORS["ink"])
paragraph("", after=3)

add_section_title("十三、最适合这个系列的写法")
add_note(
    "写作方向",
    "档案里不要大声说恐怖。只要让读者看到：记录很完整，结论很克制，关键内容被封存，幸存者不能看影像，对外说法和内部记录不一样。这就够了。",
    "F6F1E7",
)

doc.save(OUT)
print(OUT)
