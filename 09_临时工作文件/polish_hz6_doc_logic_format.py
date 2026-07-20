from datetime import datetime
from pathlib import Path
import shutil
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


PICS = Path.home() / "Pictures"
DOCX = next(
    p for p in PICS.glob("1952_HZ-6*.docx")
    if "\u6837\u672c\u7ebf\u4efb\u52a1\u62a5\u544a" in p.name and "backup" not in p.name
)
BACKUP = DOCX.with_name(DOCX.stem + "_backup_before_logic_format_" + datetime.now().strftime("%Y%m%d-%H%M%S") + DOCX.suffix)

EAST_FONT = "SimSun"
TITLE_EAST_FONT = "Microsoft YaHei"
LATIN_FONT = "Times New Roman"
MONO_FONT = "Courier New"
INK = RGBColor(28, 28, 28)
MUTED = RGBColor(80, 80, 80)


REPLACEMENTS = {
    "内容限于 HZ-6 样本任务的人员、装备、通联、影像材料和人员损失。":
        "内容限于 HZ-6 采集与回收任务的人员、装备、通联、影像材料和人员损失。",
    "克莱恩在样本点北侧地表霜晶。气体检测未见异常。":
        "克莱恩在样本点北侧采集白壳碎片与地表霜晶样品。气体检测未见异常。",
    "目标贴近地面移动，背部较长，肩部位置偏高。闪光下可见前肢比例异常，但底片没有给出完整轮廓。它经过根板空隙时没有明显停顿，对低林有一定熟悉。":
        "目标贴近地面移动，背部较长，肩部位置偏高。闪光下可见前肢比例异常，但底片没有给出完整轮廓。它经过根板空隙时没有明显停顿，像是熟悉那片低林。",
    "保密处理":
        "原始补充询问记录另行封存，本件仅保留心理评估结论。",
    "十二、评估与意见":
        "十二、处置意见",
    "Z-6B 至 HZ-6C 低林区域暂停非必要任务，直至路线、通讯和照明规程复核完成。":
        "HZ-6B 至 HZ-6C 低林区域暂停非必要任务，直至路线、通讯和照明规程复核完成。",
    "对外说失事人员以再南极洲冰层科考活动中偶遇突发性风暴丧生。":
        "对外说明：南极科研遇到突发风暴。",
}


def set_rfonts(run, east=EAST_FONT, latin=LATIN_FONT):
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east)
    r_fonts.set(qn("w:cs"), latin)


def set_style_font(style, east=EAST_FONT, latin=LATIN_FONT, size=10.5, bold=False, color=INK):
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east)
    r_fonts.set(qn("w:cs"), latin)


def replace_paragraph_text(paragraph, text):
    old_runs = list(paragraph.runs)
    bold = any(run.bold for run in old_runs)
    size = None
    if old_runs and old_runs[0].font.size:
        size = old_runs[0].font.size.pt
    paragraph.clear()
    run = paragraph.add_run(text)
    set_rfonts(run, east=TITLE_EAST_FONT if paragraph.style.name == "Heading 1" else EAST_FONT)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = INK


def polish_paragraph(paragraph):
    text = paragraph.text
    if not text:
        return False
    new = text
    for old, replacement in REPLACEMENTS.items():
        new = new.replace(old, replacement)
    if new != text:
        replace_paragraph_text(paragraph, new)
        return True
    return False


def walk_tables(tables):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                yield cell.paragraphs
                yield from walk_tables(cell.tables)


def normalize_paragraph_fonts(paragraph):
    if not paragraph.text.strip():
        return
    if paragraph.style.name == "Heading 1":
        for run in paragraph.runs:
            if run.text:
                set_rfonts(run, east=TITLE_EAST_FONT)
                run.font.size = Pt(14)
                run.bold = True
    elif paragraph.text.strip() == "HZ-6\u91c7\u96c6\u4e0e\u56de\u6536\u4efb\u52a1\u62a5\u544a":
        for run in paragraph.runs:
            if run.text:
                set_rfonts(run, east=TITLE_EAST_FONT)
                run.font.size = Pt(17.5)
                run.bold = True
    else:
        for run in paragraph.runs:
            if run.text:
                latin = MONO_FONT if any(c.isdigit() for c in run.text) and run.font.name == MONO_FONT else LATIN_FONT
                set_rfonts(run, east=EAST_FONT, latin=latin)
                if run.font.size is None:
                    run.font.size = Pt(10.5)
                run.font.color.rgb = INK

    paragraph.paragraph_format.line_spacing = 1.15
    if paragraph.style.name == "Heading 1":
        paragraph.paragraph_format.space_before = Pt(14)
        paragraph.paragraph_format.space_after = Pt(8)
    else:
        paragraph.paragraph_format.space_after = Pt(6)


def main():
    shutil.copy2(DOCX, BACKUP)
    doc = Document(DOCX)
    changed = 0

    set_style_font(doc.styles["Normal"], size=10.5)
    if "Heading 1" in doc.styles:
        set_style_font(doc.styles["Heading 1"], east=TITLE_EAST_FONT, size=14, bold=True)
    if "List Number" in doc.styles:
        set_style_font(doc.styles["List Number"], size=10.5)

    for paragraph in doc.paragraphs:
        changed += int(polish_paragraph(paragraph))
        normalize_paragraph_fonts(paragraph)

    for paragraph_group in walk_tables(doc.tables):
        for paragraph in paragraph_group:
            changed += int(polish_paragraph(paragraph))
            normalize_paragraph_fonts(paragraph)

    doc.core_properties.title = "HZ-6\u91c7\u96c6\u4e0e\u56de\u6536\u4efb\u52a1\u62a5\u544a"
    doc.core_properties.subject = "HZ-6\u91c7\u96c6\u4e0e\u56de\u6536\u4efb\u52a1\u62a5\u544a"
    doc.save(DOCX)

    with zipfile.ZipFile(DOCX) as z:
        text = ""
        for name in z.namelist():
            if name.startswith(("word/", "docProps/")) and name.endswith((".xml", ".rels")):
                text += z.read(name).decode("utf-8", errors="ignore")
        residuals = {
            "old_sample_line": text.count("\u6837\u672c\u7ebf"),
            "dropped_hz": text.count(">Z-6B"),
            "bad_external_sentence": text.count("\u4ee5\u518d\u5357\u6781\u6d32"),
        }
    print("edited", DOCX)
    print("backup", BACKUP)
    print("changed", changed)
    print("residuals", residuals)


if __name__ == "__main__":
    main()
